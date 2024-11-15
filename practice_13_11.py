import docx


doc = docx.Document('1.docx')

print(doc.paragraphs)

for i in doc.paragraphs:
    print(i.text)